#!/usr/bin/env python3
"""
Generate a formatted Word (.docx) document from an atlas JSON file.

Supports both datasets in this repo:
  - work-atlas/businesses.json  (sector -> group -> type -> examples)
  - niche-atlas/niches.json      (domain -> category -> subcategory -> niche -> micro)

Usage:
    pip install python-docx
    python make_docx.py                          # work-atlas -> work-atlas/Work-Business-Atlas.docx
    python make_docx.py --dataset niche          # niche-atlas -> niche-atlas/Niche-Atlas.docx
    python make_docx.py --input path.json --output out.docx --kind business|niche

The doc has a title page with live counts, a heading hierarchy you can navigate
with Word's Navigation Pane, and bullet lists of the leaf entries.
"""

from __future__ import annotations

import os
import json
import argparse
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_title_page(doc: Document, meta: dict, counts: dict, kind: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meta.get("name", "Atlas"))
    run.bold = True
    run.font.size = Pt(30)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run(meta.get("description", ""))
    srun.italic = True
    srun.font.size = Pt(11)
    srun.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line = f"Version {meta.get('version', '?')}  |  Generated {datetime.utcnow():%Y-%m-%d}"
    info.add_run(line).font.size = Pt(10)

    doc.add_paragraph()
    stats = doc.add_paragraph()
    stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if kind == "business":
        text = (f"{counts['l1']} sectors  -  {counts['l2']} groups  -  "
                f"{counts['l3']} business types  -  {counts['l4']} specializations\n"
                f"{counts['l3'] + counts['l4']} total sellable targets")
    else:
        text = (f"{counts['l1']} domains  -  {counts['l2']} categories  -  "
                f"{counts['l3']} subcategories  -  {counts['l4']} niches  -  "
                f"{counts['l5']} micro-niches")
    r = stats.add_run(text)
    r.bold = True
    r.font.size = Pt(12)

    doc.add_page_break()


def _add_examples(doc: Document, label: str, items: list[str]) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(label)
    run.bold = True
    if items:
        rest = p.add_run("  -  " + ", ".join(items))
        rest.italic = True
        rest.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def build_business(doc: Document, data: dict) -> dict:
    counts = {"l1": 0, "l2": 0, "l3": 0, "l4": 0, "l5": 0}
    for sector in data["sectors"]:
        counts["l1"] += 1
        doc.add_heading(sector["name"], level=1)
        for group in sector["groups"]:
            counts["l2"] += 1
            doc.add_heading(group["name"], level=2)
            for t in group["types"]:
                counts["l3"] += 1
                ex = t.get("examples", [])
                counts["l4"] += len(ex)
                _add_examples(doc, t["name"], ex)
    return counts


def build_niche(doc: Document, data: dict) -> dict:
    counts = {"l1": 0, "l2": 0, "l3": 0, "l4": 0, "l5": 0}
    for domain in data["domains"]:
        counts["l1"] += 1
        doc.add_heading(domain["name"], level=1)
        for cat in domain["categories"]:
            counts["l2"] += 1
            doc.add_heading(cat["name"], level=2)
            for sub in cat["subcategories"]:
                counts["l3"] += 1
                doc.add_heading(sub["name"], level=3)
                for niche in sub["niches"]:
                    counts["l4"] += 1
                    micro = niche.get("micro", [])
                    counts["l5"] += len(micro)
                    _add_examples(doc, niche["name"], micro)
    return counts


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--dataset", choices=["business", "niche"], default="business")
    ap.add_argument("--input", default=None)
    ap.add_argument("--output", default=None)
    ap.add_argument("--kind", choices=["business", "niche"], default=None)
    args = ap.parse_args()

    here = os.path.dirname(os.path.abspath(__file__))
    repo = os.path.dirname(here)

    if args.dataset == "business":
        inp = args.input or os.path.join(repo, "work-atlas", "businesses.json")
        out = args.output or os.path.join(repo, "work-atlas", "Work-Business-Atlas.docx")
        kind = args.kind or "business"
    else:
        inp = args.input or os.path.join(repo, "niche-atlas", "niches.json")
        out = args.output or os.path.join(repo, "niche-atlas", "Niche-Atlas.docx")
        kind = args.kind or "niche"

    with open(inp, encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    # placeholder title page added after we know counts -> build body into a temp, then assemble
    if kind == "business":
        counts = build_business(doc, data)
    else:
        counts = build_niche(doc, data)

    # Insert the title page at the very front by building a second doc.
    final = Document()
    _add_title_page(final, data.get("meta", {}), counts, kind)
    # Append body from `doc` by moving its XML body elements into `final`.
    body = final.element.body
    for child in list(doc.element.body):
        body.append(child)

    final.save(out)
    total = counts["l3"] + counts["l4"] if kind == "business" else counts["l4"] + counts["l5"]
    print(f"Wrote {out}")
    print(f"  levels: {counts}")
    print(f"  leaf entries: {total}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
